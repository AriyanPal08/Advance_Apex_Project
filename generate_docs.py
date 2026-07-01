import json
from docx import Document

with open('metrics.json', 'r') as f:
    metrics = json.load(f)

# 1. Generate Project Report
doc = Document()
doc.add_heading('Project Report - BFSI Compliance Automation Pipeline', 0)

doc.add_heading('Section 1: Problem Background', level=1)
doc.add_paragraph(
    "Financial institutions in regulated environments need to maintain continuous compliance. "
    "Compliance teams review thousands of records to identify high-risk customers and compliance gaps. "
    "The CFPB Consumer Complaint Database mirrors these records closely, containing free-text analyst narratives, "
    "inconsistent product labels, missing fields, and date formats. "
    "This project builds a data pipeline that ingests raw complaints, resolves quality issues, standardizes classification fields, "
    "and produces an analysis-ready table. "
    "This project does not build any ML model."
)

doc.add_heading('Section 2: Dataset Description', level=1)
doc.add_paragraph(
    "Dataset: CFPB Consumer Complaint Database\n"
    "Source: consumerfinance.gov/data-research/consumer-complaints\n"
    f"Structure: {metrics['before_rows']} rows and {metrics['before_cols']} columns.\n"
    "Known quality issues: missing ZIP codes and States, unstandardized text, "
    "and complaints with unresolved resolution times."
)

doc.add_heading('Section 3: Pipeline Design', level=1)
doc.add_paragraph(
    "1. Data Ingestion: Loaded dataset and preserved a before snapshot.\n"
    "2. Duplicate Removal: Dropped duplicates based on the 'Complaint ID' column to ensure uniqueness.\n"
    "3. Date Standardization: Parsed 'Date received' and 'Date sent to company' into proper datetime formats.\n"
    "4. Feature Engineering: Derived 'resolution_days' and mapped negative values to NaN to correct impossible timelines.\n"
    "5. Standardization: Stripped whitespace and title-cased 'Product' and 'Issue' to prevent mismatched aggregations.\n"
    "6. Missing Values: Handled nulls across State, ZIP code, and response fields to avoid dropping massive amounts of data.\n"
    "7. Encoding: Mapped 'Timely response?' to numeric values (1, 0, -1) and label-encoded 'Product' for analysis.\n"
    "8. Normalization: MinMax scaled 'resolution_days'."
)

doc.add_heading('Section 4: Before / After Quality Report', level=1)
table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Before Cleaning'
hdr[2].text = 'After Cleaning'
hdr[3].text = 'Improvement'

metrics_data = [
    ('Row Count', str(metrics['before_rows']), str(metrics['after_rows']), str(metrics['before_rows'] - metrics['after_rows'])),
    ('Total Nulls', str(metrics['before_nulls']), str(metrics['after_nulls']), str(metrics['before_nulls'] - metrics['after_nulls'])),
    ('Duplicate IDs', str(metrics['before_duplicates']), str(metrics['after_duplicates']), str(metrics['before_duplicates'] - metrics['after_duplicates']))
]

for item in metrics_data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

doc.add_heading('Section 5: Key Findings', level=1)
top_prod_str = list(metrics['top_products'].keys())[0]
top_prod_val = list(metrics['top_products'].values())[0]
top_state_str = list(metrics['top_states'].keys())[0]
top_state_val = list(metrics['top_states'].values())[0]

doc.add_paragraph(
    "1. Top Product Volume: The dataset reveals that the product with the most complaints is "
    f"'{top_prod_str}' with a volume of {top_prod_val:,} complaints.\n\n"
    "2. State Distribution: Geographically, the state of "
    f"'{top_state_str}' leads in complaint volume, lodging {top_state_val:,} total records.\n\n"
    "3. Timely Response Rate: Analyzing company behavior, companies that close complaints "
    "with non-monetary relief exhibit the highest timely response behavior compared to other closure methods."
)

doc.add_heading('Section 6: Limitations and Conclusion', level=1)
doc.add_paragraph(
    "Limitations: The current CFPB schema no longer contains the 'Consumer disputed?' column. "
    "Therefore, this project adapts the analysis to focus on complaint volume and timely response metrics instead. "
    "Additionally, missing narratives are simply flagged rather than utilizing NLP to extract insights. "
    "Conclusion: This project successfully processed a substantial subset of CFPB complaints, "
    "cleaning and normalizing the data to deliver a robust dataset ready for compliance trend analysis."
)

doc.save('Project_Report.docx')

# 2. Generate GenAI Prompt Log
log_doc = Document()
log_doc.add_heading('GenAI Prompt Log', 0)

log_doc.add_paragraph(
    "Student Name: Ariyan Pal\n"
    "Project: P5 — BFSI Compliance Automation Pipeline\n"
    "Dataset: CFPB Consumer Complaint Database\n"
)

table_log = log_doc.add_table(rows=1, cols=3)
hdr_log = table_log.rows[0].cells
hdr_log[0].text = 'Tool Used'
hdr_log[1].text = 'Prompt (paraphrased)'
hdr_log[2].text = 'Usage'

prompts = [
    ("ChatGPT", "How do I filter a pandas dataframe by year for an extremely large CSV without memory errors?", "Used the suggested pd.read_csv(chunksize=...) approach to build the filter_data script."),
    ("ChatGPT", "My resolution days have negative values. What is the best way to handle this?", "Used the advice to set negative values to NaN instead of dropping rows to preserve dataset integrity."),
    ("ChatGPT", "Can you show me a matplotlib template to plot the top 10 items in a bar chart cleanly?", "Adapted the horizontal bar chart formatting for the top products analysis."),
    ("ChatGPT", "The Consumer disputed column is missing from my CFPB dataset, what should I do?", "Used the advice to focus on timely response rates and volume instead and documented the schema change in limitations.")
]

for item in prompts:
    row = table_log.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]

log_doc.save('GenAI_Prompt_Log.docx')
print("Documents generated successfully.")
